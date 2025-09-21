import pymupdf
import docx
import mammoth
import io
import mimetypes
from typing import Dict, Optional, Tuple
import difflib
from dataclasses import dataclass


@dataclass
class DocumentChange:
    change_type: str  # 'added', 'modified', 'deleted'
    old_content: Optional[str]
    new_content: Optional[str]
    diff: str
    similarity_score: float


class DocumentParser:
    """Service for parsing various document formats and extracting text"""

    SUPPORTED_FORMATS = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/msword': 'doc',
        'text/plain': 'txt',
        'text/markdown': 'md',
        'application/json': 'json',
        'text/csv': 'csv'
    }

    def __init__(self):
        pass

    def extract_text(self, file_content: bytes, mime_type: str, filename: str = "") -> str:
        """Extract text from various document formats"""
        try:
            if mime_type == 'application/pdf':
                return self._extract_pdf_text(file_content)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                return self._extract_docx_text(file_content)
            elif mime_type == 'application/msword':
                return self._extract_doc_text(file_content)
            elif mime_type in ['text/plain', 'text/markdown']:
                return file_content.decode('utf-8')
            elif mime_type == 'application/json':
                import json
                return json.dumps(json.loads(file_content.decode('utf-8')), indent=2)
            elif mime_type == 'text/csv':
                return file_content.decode('utf-8')
            else:
                # Fallback to UTF-8 decoding
                return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            raise ValueError(f"Failed to extract text from {filename}: {str(e)}")

    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF using PyMuPDF"""
        doc = pymupdf.open(stream=file_content, filetype="pdf")
        text = ""

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            page_text = page.get_text()
            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"

        doc.close()
        return text.strip()

    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX using mammoth for better formatting"""
        try:
            # Use mammoth for better text extraction
            result = mammoth.extract_raw_text(io.BytesIO(file_content))
            return result.value
        except Exception:
            # Fallback to python-docx
            doc = docx.Document(io.BytesIO(file_content))
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text)

    def _extract_doc_text(self, file_content: bytes) -> str:
        """Extract text from DOC files - requires additional libraries"""
        # Note: This would require antiword or similar tools
        raise NotImplementedError("currently DOC format is not supported")

    def detect_mime_type(self, filename: str, file_content: bytes) -> str:
        """Detect MIME type from filename and content"""
        # First try by filename
        mime_type, _ = mimetypes.guess_type(filename)

        if mime_type:
            return mime_type

        # Fallback detection by content
        if file_content.startswith(b'%PDF'):
            return 'application/pdf'
        elif file_content.startswith(b'PK\x03\x04'):
            # Could be DOCX or other ZIP-based format
            if filename.lower().endswith('.docx'):
                return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        return 'text/plain'  # Default fallback

    def compare_documents(self, old_content: str, new_content: str) -> DocumentChange:
        """Compare two document contents and generate change information"""
        if not old_content:
            return DocumentChange(
                change_type='added',
                old_content=None,
                new_content=new_content,
                diff=f"+++ ADDED CONTENT +++\n{new_content}",
                similarity_score=0.0
            )

        if not new_content:
            return DocumentChange(
                change_type='deleted',
                old_content=old_content,
                new_content=None,
                diff=f"--- DELETED CONTENT ---\n{old_content}",
                similarity_score=0.0
            )

        # Calculate similarity
        similarity = difflib.SequenceMatcher(None, old_content, new_content).ratio()

        # Generate unified diff
        diff = '\n'.join(difflib.unified_diff(
            old_content.splitlines(keepends=True),
            new_content.splitlines(keepends=True),
            fromfile='old_version',
            tofile='new_version',
            lineterm=''
        ))

        return DocumentChange(
            change_type='modified',
            old_content=old_content,
            new_content=new_content,
            diff=diff,
            similarity_score=similarity
        )

    def extract_metadata(self, file_content: bytes, mime_type: str) -> Dict:
        """Extract metadata from documents"""
        metadata = {
            'file_size': len(file_content),
            'mime_type': mime_type
        }

        try:
            if mime_type == 'application/pdf':
                doc = pymupdf.open(stream=file_content, filetype="pdf")
                metadata.update({
                    'page_count': len(doc),
                    'title': doc.metadata.get('title', ''),
                    'author': doc.metadata.get('author', ''),
                    'subject': doc.metadata.get('subject', ''),
                    'creator': doc.metadata.get('creator', ''),
                    'creation_date': doc.metadata.get('creationDate', ''),
                    'modification_date': doc.metadata.get('modDate', '')
                })
                doc.close()

            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                doc = docx.Document(io.BytesIO(file_content))
                props = doc.core_properties
                metadata.update({
                    'title': props.title or '',
                    'author': props.author or '',
                    'subject': props.subject or '',
                    'created': props.created.isoformat() if props.created else '',
                    'modified': props.modified.isoformat() if props.modified else '',
                    'last_modified_by': props.last_modified_by or ''
                })

        except Exception as e:
            metadata['metadata_error'] = str(e)

        return metadata


# Usage example and integration with existing models
class DocumentVersionService:
    """Service for managing document versions and changes"""

    def __init__(self, parser: DocumentParser):
        self.parser = parser

    def process_file_upload(self, file_content: bytes, filename: str,
                            repository_id: str, old_file_content: str = None) -> Dict:
        """Process uploaded file and generate change information"""
        # Detect MIME type
        mime_type = self.parser.detect_mime_type(filename, file_content)

        # Extract text content
        text_content = self.parser.extract_text(file_content, mime_type, filename)

        # Compare with old version if exists
        change_info = None
        if old_file_content is not None:
            change_info = self.parser.compare_documents(old_file_content, text_content)

        # Extract metadata
        metadata = self.parser.extract_metadata(file_content, mime_type)

        return {
            'text_content': text_content,
            'mime_type': mime_type,
            'metadata': metadata,
            'change_info': change_info,
            'file_size': len(file_content)
        }